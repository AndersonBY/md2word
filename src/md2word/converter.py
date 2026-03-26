"""Public converter facade and orchestration."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import markdown2
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from html4docx import HtmlToDocx

from .config import Config
from .conversion import markdown as markdown_compat
from .conversion.blocks import (
    add_code_block_to_document,
    extract_blockquotes,
    extract_code_blocks,
    replace_blockquote_placeholders,
    replace_code_block_placeholders,
    style_inline_code_in_document,
)
from .conversion.common import hex_to_rgb, print_error, print_info
from .conversion.images import (
    decode_data_uri_image,
    download_image,
    ensure_local_image_compatible,
    filter_unrecognized_images,
    is_docx_image_supported,
    process_image_content,
    process_markdown_images,
    resize_images_in_document,
    sanitize_html_images,
)
from .conversion.numbering import CHINESE_NUMBERS, HeadingNumbering, number_to_chinese
from .conversion.styles import (
    apply_style_to_paragraph,
    apply_style_to_run,
    apply_styles_to_document,
    apply_table_styles,
    get_heading_level,
    is_code_block_paragraph,
)
from .conversion.toc import add_toc
from .latex import extract_latex_formulas, replace_formula_placeholders

_MARKDOWN2_PUNCTUATED_EMPHASIS_COMPAT = markdown_compat._MARKDOWN2_PUNCTUATED_EMPHASIS_COMPAT


@dataclass(slots=True)
class PreparedMarkdown:
    """Intermediate data produced before Word rendering."""

    html_content: str
    formulas: list[Any]
    code_blocks: list[dict[str, str]]
    inline_codes: list[str]
    blockquotes: list[str]


def fix_markdown2_punctuated_emphasis_html(html_content: str) -> str:
    """Compatibility wrapper that preserves the legacy monkeypatch surface."""
    previous = markdown_compat._MARKDOWN2_PUNCTUATED_EMPHASIS_COMPAT
    markdown_compat._MARKDOWN2_PUNCTUATED_EMPHASIS_COMPAT = _MARKDOWN2_PUNCTUATED_EMPHASIS_COMPAT
    try:
        return markdown_compat.fix_markdown2_punctuated_emphasis_html(html_content)
    finally:
        markdown_compat._MARKDOWN2_PUNCTUATED_EMPHASIS_COMPAT = previous


def _prepare_markdown(markdown_content: str, config: Config) -> PreparedMarkdown:
    """Prepare markdown for document rendering."""
    processed_content, formulas = extract_latex_formulas(markdown_content)
    if formulas:
        print_info(f"Detected {len(formulas)} LaTeX formulas")

    processed_content = process_markdown_images(processed_content, config)
    html_content = markdown2.markdown(
        processed_content,
        extras=["tables", "cuddled-lists", "fenced-code-blocks", "header-ids"],
    )
    html_content = fix_markdown2_punctuated_emphasis_html(html_content)
    html_content = sanitize_html_images(html_content, config)

    html_content, code_blocks, inline_codes = extract_code_blocks(html_content)
    if code_blocks:
        print_info(f"Extracted {len(code_blocks)} code blocks")
    if inline_codes:
        print_info(f"Found {len(inline_codes)} inline code snippets")

    html_content, blockquotes = extract_blockquotes(html_content)
    if blockquotes:
        print_info(f"Extracted {len(blockquotes)} blockquotes")

    return PreparedMarkdown(
        html_content=html_content,
        formulas=formulas,
        code_blocks=code_blocks,
        inline_codes=inline_codes,
        blockquotes=blockquotes,
    )


def _render_html_to_document(html_content: str):
    """Render prepared HTML into a Word document with image fallbacks."""
    document = Document()
    parser = HtmlToDocx()

    try:
        parser.add_html_to_document(html_content, document)
    except UnrecognizedImageError as e:
        print_error(f"UnrecognizedImageError, retrying without problematic images: {e}")
        html_filtered = filter_unrecognized_images(html_content)
        document = Document()
        parser = HtmlToDocx()
        try:
            parser.add_html_to_document(html_filtered, document)
        except UnrecognizedImageError as e2:
            print_error(f"Still failing, removing all images: {e2}")
            html_without_images = re.sub(r"<img[^>]*>", "", html_filtered, flags=re.IGNORECASE)
            document = Document()
            HtmlToDocx().add_html_to_document(html_without_images, document)

    return document


def _finalize_document(
    document,
    prepared: PreparedMarkdown,
    config: Config,
    toc: bool,
    toc_title: str,
    toc_max_level: int,
) -> None:
    """Apply post-render document transformations."""
    if prepared.code_blocks:
        replace_code_block_placeholders(document, prepared.code_blocks, config)
    if prepared.blockquotes:
        replace_blockquote_placeholders(document, prepared.blockquotes, config)
    if prepared.formulas:
        replace_formula_placeholders(document, prepared.formulas)

    apply_styles_to_document(document, config)

    if prepared.inline_codes:
        style_inline_code_in_document(document, config)

    resize_images_in_document(document, config.max_image_width_inches)

    if toc and document.paragraphs:
        add_toc(document, title=toc_title, max_level=toc_max_level)


def convert(
    markdown_content: str,
    output_path: str | Path,
    config: Config | None = None,
    toc: bool = False,
    toc_title: str = "目录",
    toc_max_level: int = 3,
) -> Path:
    """
    Convert Markdown content to a Word document.

    Args:
        markdown_content: Markdown text content
        output_path: Output file path
        config: Configuration object (uses defaults if None)
        toc: Whether to add table of contents
        toc_title: TOC title
        toc_max_level: Maximum heading level for TOC

    Returns:
        Path to the output file
    """
    config = Config() if config is None else config
    output_path = Path(output_path)

    prepared = _prepare_markdown(markdown_content, config)
    document = _render_html_to_document(prepared.html_content)
    _finalize_document(document, prepared, config, toc, toc_title, toc_max_level)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print_info(f"Document saved: {output_path}")
    return output_path


def convert_file(
    input_path: str | Path,
    output_path: str | Path | None = None,
    config: Config | str | Path | None = None,
    toc: bool = False,
    toc_title: str = "目录",
    toc_max_level: int = 3,
) -> Path:
    """
    Convert a Markdown file to a Word document.

    Args:
        input_path: Input Markdown file path
        output_path: Output file path (defaults to input with .docx extension)
        config: Configuration object or path to config file
        toc: Whether to add table of contents
        toc_title: TOC title
        toc_max_level: Maximum heading level for TOC

    Returns:
        Path to the output file
    """
    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    output_path = input_path.with_suffix(".docx") if output_path is None else Path(output_path)

    if config is None:
        config = Config()
    elif isinstance(config, str | Path):
        config = Config.from_file(config)

    markdown_content = input_path.read_text(encoding="utf-8")
    if markdown_content.startswith("```markdown") and markdown_content.endswith("```"):
        markdown_content = markdown_content[12:-3]

    return convert(
        markdown_content,
        output_path,
        config,
        toc=toc,
        toc_title=toc_title,
        toc_max_level=toc_max_level,
    )


__all__ = [
    "CHINESE_NUMBERS",
    "HeadingNumbering",
    "_MARKDOWN2_PUNCTUATED_EMPHASIS_COMPAT",
    "add_code_block_to_document",
    "add_toc",
    "apply_style_to_paragraph",
    "apply_style_to_run",
    "apply_styles_to_document",
    "apply_table_styles",
    "convert",
    "convert_file",
    "decode_data_uri_image",
    "download_image",
    "ensure_local_image_compatible",
    "extract_blockquotes",
    "extract_code_blocks",
    "filter_unrecognized_images",
    "fix_markdown2_punctuated_emphasis_html",
    "get_heading_level",
    "hex_to_rgb",
    "is_code_block_paragraph",
    "is_docx_image_supported",
    "number_to_chinese",
    "print_error",
    "print_info",
    "process_image_content",
    "process_markdown_images",
    "replace_blockquote_placeholders",
    "replace_code_block_placeholders",
    "resize_images_in_document",
    "sanitize_html_images",
    "style_inline_code_in_document",
]
