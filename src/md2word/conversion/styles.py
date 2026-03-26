"""Document styling helpers."""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..config import Config, StyleConfig
from .common import hex_to_rgb
from .numbering import HeadingNumbering


def apply_style_to_run(run, style_config: StyleConfig) -> None:
    """Apply style configuration to a run."""
    run.font.name = style_config.font_name
    run.font.size = Pt(style_config.font_size)
    run.font.bold = run.font.bold or style_config.bold
    run.font.italic = run.font.italic or style_config.italic

    r, g, b = hex_to_rgb(style_config.color)
    run.font.color.rgb = RGBColor(r, g, b)

    if run._element.rPr is not None:
        r_fonts = run._element.rPr.rFonts
        if r_fonts is not None:
            r_fonts.set(qn("w:eastAsia"), style_config.font_name)


def apply_style_to_paragraph(paragraph, style_config: StyleConfig) -> None:
    """Apply style configuration to a paragraph."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(style_config.space_before)
    paragraph_format.space_after = Pt(style_config.space_after)

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if style_config.alignment in alignment_map:
        paragraph_format.alignment = alignment_map[style_config.alignment]

    if style_config.line_spacing_rule == "exact" and style_config.line_spacing_value:
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(style_config.line_spacing_value)
    elif style_config.line_spacing_rule == "at_least" and style_config.line_spacing_value:
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        paragraph_format.line_spacing = Pt(style_config.line_spacing_value)
    elif style_config.line_spacing_rule == "single":
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif style_config.line_spacing_rule == "1.5":
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif style_config.line_spacing_rule == "double":
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif style_config.line_spacing_rule == "multiple":
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        if style_config.line_spacing_value:
            paragraph_format.line_spacing = style_config.line_spacing_value
        elif style_config.line_spacing > 0:
            paragraph_format.line_spacing = style_config.line_spacing
    elif style_config.line_spacing > 0:
        paragraph_format.line_spacing = style_config.line_spacing

    if style_config.left_indent > 0:
        paragraph_format.left_indent = Inches(style_config.left_indent)
    if style_config.first_line_indent > 0:
        indent_pt = style_config.first_line_indent * style_config.font_size
        paragraph_format.first_line_indent = Pt(indent_pt)


def get_heading_level(paragraph) -> int | None:
    """Get heading level of paragraph, or ``None`` if it is not a heading."""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.replace("Heading ", "").replace("Heading", ""))
        except ValueError:
            return None
    return None


def is_code_block_paragraph(paragraph) -> bool:
    """Check whether a paragraph is a code block."""
    paragraph_properties = paragraph._element.pPr
    if paragraph_properties is not None:
        shading = paragraph_properties.find(qn("w:shd"))
        if shading is not None:
            fill = shading.get(qn("w:fill"))
            if fill and fill.lower() not in ("auto", "ffffff", "none"):
                return True
    return False


def apply_styles_to_document(document, config: Config) -> None:
    """Apply style configuration to a document."""
    numbering = HeadingNumbering()

    for paragraph in document.paragraphs:
        if is_code_block_paragraph(paragraph):
            continue

        heading_level = get_heading_level(paragraph)
        if heading_level is not None:
            style_name = f"heading_{heading_level}"
            style_config = config.get_style(style_name)
            if style_config.numbering_format and paragraph.runs:
                number_text = numbering.get_number(heading_level, style_config.numbering_format)
                if number_text:
                    first_run = paragraph.runs[0]
                    first_run.text = number_text + first_run.text
        else:
            style_config = config.get_style("body")

        apply_style_to_paragraph(paragraph, style_config)
        for run in paragraph.runs:
            apply_style_to_run(run, style_config)

    apply_table_styles(document, config)


def apply_table_styles(document, config: Config) -> None:
    """Apply table styling from configuration."""
    table_config = config.table
    border_style_map = {
        "single": "single",
        "double": "double",
        "dotted": "dotted",
        "dashed": "dashed",
        "none": "nil",
    }
    border_val = border_style_map.get(table_config.border_style, "single")

    for table in document.tables:
        if table_config.width_mode == "full":
            table.autofit = False
            table.allow_autofit = False
            table_properties = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement("w:tblPr")
            table_width = OxmlElement("w:tblW")
            table_width.set(qn("w:w"), "5000")
            table_width.set(qn("w:type"), "pct")
            table_properties.append(table_width)
            if table._tbl.tblPr is None:
                table._tbl.insert(0, table_properties)
        elif table_config.width_mode == "fixed" and table_config.width_inches:
            table.autofit = False
            table_properties = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement("w:tblPr")
            table_width = OxmlElement("w:tblW")
            table_width.set(qn("w:w"), str(int(table_config.width_inches * 1440)))
            table_width.set(qn("w:type"), "dxa")
            table_properties.append(table_width)
            if table._tbl.tblPr is None:
                table._tbl.insert(0, table_properties)

        for i, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    style_config = config.get_style("table_header" if i == 0 else "table_cell")
                    apply_style_to_paragraph(paragraph, style_config)
                    for run in paragraph.runs:
                        apply_style_to_run(run, style_config)

                tc_pr = cell._tc.get_or_add_tcPr()

                if i == 0 and table_config.header_background_color:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:val"), "clear")
                    shading.set(qn("w:color"), "auto")
                    shading.set(qn("w:fill"), table_config.header_background_color)
                    tc_pr.append(shading)
                elif i > 0:
                    if table_config.alternating_row_color and i % 2 == 0:
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:val"), "clear")
                        shading.set(qn("w:color"), "auto")
                        shading.set(qn("w:fill"), table_config.alternating_row_color)
                        tc_pr.append(shading)
                    elif table_config.cell_background_color:
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:val"), "clear")
                        shading.set(qn("w:color"), "auto")
                        shading.set(qn("w:fill"), table_config.cell_background_color)
                        tc_pr.append(shading)

                tc_mar = OxmlElement("w:tcMar")
                for side, value in [
                    ("top", table_config.cell_padding_top),
                    ("bottom", table_config.cell_padding_bottom),
                    ("left", table_config.cell_padding_left),
                    ("right", table_config.cell_padding_right),
                ]:
                    margin = OxmlElement(f"w:{side}")
                    margin.set(qn("w:w"), str(int(value * 20)))
                    margin.set(qn("w:type"), "dxa")
                    tc_mar.append(margin)
                tc_pr.append(tc_mar)

                if border_val != "nil":
                    tc_borders = OxmlElement("w:tcBorders")
                    for side in ["top", "left", "bottom", "right"]:
                        border = OxmlElement(f"w:{side}")
                        border.set(qn("w:val"), border_val)
                        border.set(qn("w:sz"), str(table_config.border_width))
                        border.set(qn("w:color"), table_config.border_color)
                        tc_borders.append(border)
                    tc_pr.append(tc_borders)


__all__ = [
    "apply_style_to_paragraph",
    "apply_style_to_run",
    "apply_styles_to_document",
    "apply_table_styles",
    "get_heading_level",
    "is_code_block_paragraph",
]
