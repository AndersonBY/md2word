"""Code block, blockquote, and inline-code helpers."""

from __future__ import annotations

import re
from copy import deepcopy

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..config import Config
from .common import hex_to_rgb, print_info
from .styles import apply_style_to_paragraph

_INLINE_CODE_MARKER_RE = re.compile(r"⟦CODE⟧(.*?)⟦/CODE⟧")


def extract_blockquotes(html_content: str) -> tuple[str, list[str]]:
    """Extract blockquotes from HTML and replace them with placeholders."""
    blockquotes: list[str] = []

    def save_blockquote(match: re.Match[str]) -> str:
        block_html = match.group(0)
        text = re.sub(r"<[^>]+>", "", block_html).strip()
        text = text.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
        text = text.replace("&quot;", '"').replace("&#39;", "'")

        blockquotes.append(text)
        placeholder = f"__BLOCKQUOTE_PLACEHOLDER_{len(blockquotes) - 1}__"
        return f"<p>{placeholder}</p>"

    html_content = re.sub(
        r"<blockquote[^>]*>.*?</blockquote>",
        save_blockquote,
        html_content,
        flags=re.DOTALL | re.IGNORECASE,
    )

    return html_content, blockquotes


def replace_blockquote_placeholders(document, blockquotes: list[str], config: Config) -> None:
    """Replace blockquote placeholders with styled paragraphs."""
    if not blockquotes:
        return

    style_config = config.get_style("blockquote")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for idx, quote_text in enumerate(blockquotes):
            placeholder = f"__BLOCKQUOTE_PLACEHOLDER_{idx}__"
            if text != placeholder:
                continue

            paragraph.clear()
            run = paragraph.add_run(quote_text)
            run.font.name = style_config.font_name
            run.font.size = Pt(style_config.font_size)
            run.font.italic = style_config.italic
            run.font.bold = style_config.bold

            r, g, b = hex_to_rgb(style_config.color)
            run.font.color.rgb = RGBColor(r, g, b)

            run_properties = run._element.get_or_add_rPr()
            run_fonts = run_properties.get_or_add_rFonts()
            run_fonts.set(qn("w:eastAsia"), style_config.font_name)

            apply_style_to_paragraph(paragraph, style_config)

            paragraph_properties = paragraph._element.get_or_add_pPr()
            paragraph_border = OxmlElement("w:pBdr")
            left_border = OxmlElement("w:left")
            left_border.set(qn("w:val"), "single")
            left_border.set(qn("w:sz"), "24")
            left_border.set(qn("w:space"), "4")
            left_border.set(qn("w:color"), style_config.color)
            paragraph_border.append(left_border)
            paragraph_properties.append(paragraph_border)

            print_info(f"Styled blockquote ({len(quote_text)} chars)")
            break


def extract_code_blocks(html_content: str) -> tuple[str, list[dict[str, str]], list[str]]:
    """Extract code blocks from HTML and replace them with placeholders."""
    code_blocks: list[dict[str, str]] = []
    inline_codes: list[str] = []

    def save_code_block(match: re.Match[str]) -> str:
        block_html = match.group(0)
        clean_content = re.sub(r"<span[^>]*>", "", block_html)
        clean_content = re.sub(r"</span>", "", clean_content)

        code_match = re.search(r"<code[^>]*>(.*?)</code>", clean_content, flags=re.DOTALL | re.IGNORECASE)
        if code_match:
            code_text = code_match.group(1)
            code_text = code_text.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
            code_text = code_text.replace("&quot;", '"').replace("&#39;", "'")
        else:
            pre_match = re.search(r"<pre[^>]*>(.*?)</pre>", clean_content, flags=re.DOTALL | re.IGNORECASE)
            code_text = pre_match.group(1) if pre_match else ""

        placeholder = f"__CODE_BLOCK_PLACEHOLDER_{len(code_blocks)}__"
        code_blocks.append({"code": code_text.strip(), "placeholder": placeholder})
        return f"<p>{placeholder}</p>"

    html_content = re.sub(
        r'<div[^>]*class="codehilite"[^>]*>\s*<pre[^>]*>.*?</pre>\s*</div>',
        save_code_block,
        html_content,
        flags=re.DOTALL | re.IGNORECASE,
    )
    html_content = re.sub(
        r"<pre[^>]*>.*?</pre>",
        save_code_block,
        html_content,
        flags=re.DOTALL | re.IGNORECASE,
    )

    def mark_inline_code(match: re.Match[str]) -> str:
        code_text = match.group(1)
        code_text = code_text.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
        code_text = code_text.replace("&quot;", '"').replace("&#39;", "'")
        inline_codes.append(code_text)
        return f"⟦CODE⟧{code_text}⟦/CODE⟧"

    html_content = re.sub(r"<code>([^<]*)</code>", mark_inline_code, html_content, flags=re.IGNORECASE)
    return html_content, code_blocks, inline_codes


def add_code_block_to_document(paragraph, code_text: str, config: Config) -> None:
    """Replace a placeholder paragraph with properly formatted code block."""
    code_style = config.get_style("code")
    font_name = code_style.font_name
    font_size = code_style.font_size
    bg_color = code_style.background_color or "f5f5f5"

    paragraph.clear()

    lines = code_text.split("\n")
    for i, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if i < len(lines) - 1:
            run.add_break()

    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.0

    paragraph_properties = paragraph._element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), bg_color)
    paragraph_properties.append(shading)


def replace_code_block_placeholders(document, code_blocks: list[dict[str, str]], config: Config) -> None:
    """Replace code block placeholders in the document with formatted code."""
    if not code_blocks:
        return

    placeholder_map = {block["placeholder"]: block["code"] for block in code_blocks}
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in placeholder_map:
            add_code_block_to_document(paragraph, placeholder_map[text], config)
            print_info(f"Added code block ({len(placeholder_map[text])} chars)")


def _get_run_element_text(run_element) -> str:
    """Return plain text for a Word run element."""
    parts = []
    for child in run_element:
        if child.tag == qn("w:t"):
            parts.append(child.text or "")
    return "".join(parts)


def _set_run_element_text(run_element, text: str) -> None:
    """Replace run contents with a single text node, preserving run properties."""
    for child in list(run_element):
        if child.tag != qn("w:rPr"):
            run_element.remove(child)

    text_element = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace() or "  " in text:
        text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run_element.append(text_element)


def _clone_run_element_with_text(run_element, text: str):
    """Clone a run element and replace its text content."""
    cloned = deepcopy(run_element)
    _set_run_element_text(cloned, text)
    return cloned


def _apply_inline_code_style_to_run_element(
    run_element, code_font_name: str, code_font_size: int | float, bg_color: str
) -> None:
    """Apply inline-code styling directly to a Word run element."""
    run_properties = run_element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        run_fonts.set(qn(f"w:{attr}"), code_font_name)

    size_value = str(int(float(code_font_size) * 2))
    for tag_name in ("w:sz", "w:szCs"):
        size = run_properties.find(qn(tag_name))
        if size is None:
            size = OxmlElement(tag_name)
            run_properties.append(size)
        size.set(qn("w:val"), size_value)

    existing_shading = run_properties.find(qn("w:shd"))
    if existing_shading is not None:
        run_properties.remove(existing_shading)

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), bg_color)
    run_properties.append(shading)


def _split_inline_code_segments(text: str) -> list[tuple[bool, str]]:
    """Split text into normal and inline-code segments based on markers."""
    segments: list[tuple[bool, str]] = []
    last_end = 0

    for match in _INLINE_CODE_MARKER_RE.finditer(text):
        if match.start() > last_end:
            segments.append((False, text[last_end:match.start()]))
        segments.append((True, match.group(1)))
        last_end = match.end()

    if last_end < len(text):
        segments.append((False, text[last_end:]))

    return [(is_code, value) for is_code, value in segments if value]


def _replace_inline_code_markers_in_run_element(
    run_element, code_font_name: str, code_font_size: int | float, bg_color: str
) -> bool:
    """Replace inline-code markers inside a run while preserving its container structure."""
    text = _get_run_element_text(run_element)
    if "⟦CODE⟧" not in text:
        return False

    segments = _split_inline_code_segments(text)
    if not segments:
        return False

    parent = run_element.getparent()
    insert_at = parent.index(run_element)
    new_runs = []

    for is_code, segment_text in segments:
        new_run = _clone_run_element_with_text(run_element, segment_text)
        if is_code:
            _apply_inline_code_style_to_run_element(new_run, code_font_name, code_font_size, bg_color)
        new_runs.append(new_run)

    parent.remove(run_element)
    for offset, new_run in enumerate(new_runs):
        parent.insert(insert_at + offset, new_run)

    return True


def style_inline_code_in_document(document, config: Config) -> None:
    """Find and style inline code marked with special markers."""
    code_style = config.get_style("code")
    code_font_name = code_style.font_name
    code_font_size = code_style.font_size
    bg_color = code_style.background_color or "f5f5f5"

    for paragraph in document.paragraphs:
        if "⟦CODE⟧" not in paragraph.text:
            continue

        for child in list(paragraph._p):
            if child.tag == qn("w:r"):
                _replace_inline_code_markers_in_run_element(child, code_font_name, code_font_size, bg_color)
            elif child.tag == qn("w:hyperlink"):
                for run_element in list(child):
                    if run_element.tag == qn("w:r"):
                        _replace_inline_code_markers_in_run_element(
                            run_element, code_font_name, code_font_size, bg_color
                        )


__all__ = [
    "add_code_block_to_document",
    "extract_blockquotes",
    "extract_code_blocks",
    "replace_blockquote_placeholders",
    "replace_code_block_placeholders",
    "style_inline_code_in_document",
]
