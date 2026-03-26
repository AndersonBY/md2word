"""Table-of-contents helpers."""

from __future__ import annotations

from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from .common import print_info


def add_toc(document, title: str = "目录", max_level: int = 3) -> None:
    """Add a table of contents at the beginning of the document."""
    toc_title = document.paragraphs[0].insert_paragraph_before(title)
    toc_title.style = document.styles["Heading 1"]

    toc_paragraph = toc_title.insert_paragraph_before("")
    run = toc_paragraph.add_run()

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(field_begin)

    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
    run._r.append(instruction_text)

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(field_separate)

    placeholder_run = toc_paragraph.add_run("Right-click here and select 'Update Field' to generate TOC")
    placeholder_run.italic = True
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_end)

    page_break_paragraph = toc_title.insert_paragraph_before("")
    page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)

    print_info(f"Added TOC (levels 1-{max_level})")


__all__ = ["add_toc"]
